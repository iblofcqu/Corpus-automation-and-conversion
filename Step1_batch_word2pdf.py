"""
批量将Word文档转化为PDF文件的工具
支持 .docx 和 .doc 文件转换为 PDF 格式
"""
import argparse
import os
import subprocess
import sys
from pathlib import Path


def convert_docx_to_pdf(docx_path, pdf_path):
    """
    Convert .docx file to PDF by converting to images first, which preserves formatting
    """
    try:
        import os

        import docx2pdf

        # Convert using docx2pdf - this preserves formatting and handles Unicode paths better
        abs_docx_path = os.path.abspath(str(docx_path))
        abs_pdf_path = os.path.abspath(str(pdf_path))

        # Use docx2pdf to convert (this uses the system's Word if available)
        docx2pdf.convert(abs_docx_path, abs_pdf_path)

        # Check if the PDF was created successfully
        if os.path.exists(abs_pdf_path):
            # Successfully converted
            return True
        # If it didn't create the expected file, check if it created it with .docx location
        expected_pdf_path = abs_docx_path.replace('.docx', '.pdf')
        if os.path.exists(expected_pdf_path):
            # Move the file to the intended destination
            import shutil
            shutil.move(expected_pdf_path, abs_pdf_path)
            return True

    except ImportError:
        print("错误: 需要安装 docx2pdf 库来保持Word格式")
        print("请运行: pip install docx2pdf")
        return False
    except Exception as e:
        # If docx2pdf fails, fallback to text-based conversion
        print(f"使用 docx2pdf 转换失败: {str(e)[:50]}...")
        print("使用备用方法转换（格式可能不完全保留）...")

        # Fallback implementation that creates a basic PDF
        try:
            from io import BytesIO

            from docx import Document  # This requires python-docx
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.units import inch
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.pdfgen import canvas

            # Register a Chinese font (using system font or fallback to a common Chinese font)
            chinese_font_paths = [
                # Windows fonts
                "C:/Windows/Fonts/simhei.ttf",   # SimHei (黑体)
                "C:/Windows/Fonts/simsun.ttc",   # SimSun (宋体)
                "C:/Windows/Fonts/msyh.ttc",     # Microsoft YaHei (微软雅黑)
                "C:/Windows/Fonts/msyhbd.ttc",   # Microsoft YaHei Bold
            ]

            font_path = None
            for path in chinese_font_paths:
                if os.path.exists(path):
                    font_path = path
                    break

            # Create PDF in memory
            pdf_buffer = BytesIO()
            c = canvas.Canvas(pdf_buffer, pagesize=letter)

            # Register and use Chinese font if available
            if font_path:
                font_name = "ChineseFont"
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                c.setFont(font_name, 12)  # Default font size
            else:
                # Fallback to default font (will not display Chinese properly)
                c.setFont("Helvetica", 12)

            width, height = letter

            # Set margins and starting position
            x_margin = 1 * inch  # 1 inch left/right margin
            y_margin = 1 * inch  # 1 inch top/bottom margin
            y_position = height - y_margin
            line_height = 14  # Points per line

            # Load the docx document using python-docx
            doc = Document(docx_path)

            # Process paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    # Add extra space for empty paragraphs
                    y_position -= line_height
                    continue

                # Handle the text with basic line breaking
                line_width = width - (2 * x_margin)
                current_line = ""

                for char in text:
                    test_line = current_line + char
                    # Estimation for character width
                    if font_path:
                        estimated_width = len(test_line) * 8  # Rough estimate for Chinese chars
                    else:
                        estimated_width = len(test_line) * 6  # For Latin chars

                    if estimated_width <= line_width:
                        current_line = test_line
                    else:
                        # Draw the current line
                        if y_position < y_margin + line_height:
                            c.showPage()
                            y_position = height - y_margin
                            if font_path:
                                c.setFont(font_name, 12)
                            else:
                                c.setFont("Helvetica", 12)

                        c.drawString(x_margin, y_position, current_line)
                        y_position -= line_height
                        current_line = char  # Start new line with the character that didn't fit

                # Draw any remaining text in the line
                if current_line:
                    if y_position < y_margin + line_height:
                        c.showPage()
                        y_position = height - y_margin
                        if font_path:
                            c.setFont(font_name, 12)
                        else:
                            c.setFont("Helvetica", 12)

                    c.drawString(x_margin, y_position, current_line)
                    y_position -= line_height

            # Process tables
            for table in doc.tables:
                y_position -= line_height * 2  # Add space before table
                if y_position < y_margin + (len(table.rows) * line_height):
                    c.showPage()
                    y_position = height - y_margin
                    if font_path:
                        c.setFont(font_name, 12)
                    else:
                        c.setFont("Helvetica", 12)

                # Simple table representation
                for i, row in enumerate(table.rows):
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])

                    # Break long table text similar to paragraphs
                    line_width = width - (2 * x_margin)
                    current_line = ""

                    for char in row_text:
                        test_line = current_line + char
                        if font_path:
                            estimated_width = len(test_line) * 8
                        else:
                            estimated_width = len(test_line) * 6

                        if estimated_width <= line_width:
                            current_line = test_line
                        else:
                            if y_position < y_margin + line_height:
                                c.showPage()
                                y_position = height - y_margin
                                if font_path:
                                    c.setFont(font_name, 12)
                                else:
                                    c.setFont("Helvetica", 12)

                            c.drawString(x_margin, y_position, current_line)
                            y_position -= line_height
                            current_line = char

                    # Draw any remaining text
                    if current_line:
                        if y_position < y_margin + line_height:
                            c.showPage()
                            y_position = height - y_margin
                            if font_path:
                                c.setFont(font_name, 12)
                            else:
                                c.setFont("Helvetica", 12)

                        c.drawString(x_margin, y_position, current_line)
                        y_position -= line_height

            c.save()
            pdf_bytes = pdf_buffer.getvalue()
            pdf_buffer.close()

            # Write PDF to file
            with open(pdf_path, 'wb') as f:
                f.write(pdf_bytes)

            return True

        except Exception as e:
            # Safely print error message by avoiding problematic characters
            error_msg = f"转换 .docx 文件时出错: {str(e)[:100]}..."
            safe_error = error_msg.encode('ascii', errors='ignore').decode('ascii', errors='replace')
            print(safe_error)
            return False


def convert_doc_to_pdf(doc_path, pdf_path):
    """
    Convert .doc file to PDF using system tools
    """
    try:
        import platform

        if platform.system() == "Windows":
            # Try using win32com if Microsoft Word is installed
            try:
                import win32com.client as win32
                word = win32.Dispatch('Word.Application')
                word.Visible = False

                import os
                # Open and convert the document - use absolute paths
                abs_doc_path = os.path.abspath(str(doc_path))
                abs_pdf_path = os.path.abspath(str(pdf_path))

                doc_obj = word.Documents.Open(abs_doc_path)
                doc_obj.SaveAs2(abs_pdf_path, FileFormat=17)  # 17 = PDF format
                doc_obj.Close()
                word.Quit()

                return True
            except ImportError:
                print("错误: Windows 系统需要安装 pywin32 库来处理 .doc 文件")
                print("请运行: pip install pywin32 (需先安装 Microsoft Word)")
                return False
            except Exception as e:
                # Safely print error message by avoiding problematic characters
                error_msg = f"使用 Word 转换 .doc 文件失败: {str(e)[:100]}..."
                safe_error = error_msg.encode('ascii', errors='ignore').decode('ascii', errors='replace')
                print(safe_error)
                print("尝试使用替代方法...")
        else:
            # For Linux/Mac, try antiword + pandoc or LibreOffice
            try:
                # Try LibreOffice in headless mode
                result = subprocess.run([
                    'soffice',
                    '--headless',
                    '--convert-to',
                    'pdf',
                    '--outdir',
                    str(Path(pdf_path).parent),
                    str(doc_path)
                ], capture_output=True, text=True, check=True)

                # Check if PDF was created
                expected_pdf = doc_path.with_suffix('.pdf')
                if expected_pdf.exists():
                    # Move to the target location
                    expected_pdf.rename(pdf_path)
                    return True
                print("LibreOffice 转换未生成预期的 PDF 文件")
                return False
            except (subprocess.CalledProcessError, FileNotFoundError):
                print("错误: 需要安装 LibreOffice 来处理 .doc 文件")
                print("请安装 LibreOffice 后重试")
                return False
    except Exception as e:
        # Safely print error message by avoiding problematic characters
        error_msg = f"转换 .doc 文件时出错: {str(e)[:100]}..."
        safe_error = error_msg.encode('ascii', errors='ignore').decode('ascii', errors='replace')
        print(safe_error)
        return False


def copy_pdf_to_output(pdf_path, output_path):
    """
    Copy PDF file to output directory
    """
    try:
        import shutil
        shutil.copy2(pdf_path, output_path)
        return True
    except Exception as e:
        print(f"复制 PDF 文件时出错: {e}")
        return False


def batch_convert_word_to_pdf(input_path, output_dir):
    """
    批量转换Word文档到PDF
    """
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)

    # Determine if input is a file or directory
    input_path = Path(input_path)

    if input_path.is_file():
        # Single file processing
        file_list = [input_path]
    elif input_path.is_dir():
        # Directory processing - find all .docx, .doc, and .pdf files
        file_list = list(input_path.glob('*.docx')) + \
                   list(input_path.glob('*.doc')) + \
                   list(input_path.glob('*.pdf'))
    else:
        print(f"Error: Input path does not exist - {input_path}")
        return False

    if not file_list:
        print("No supported files found (.docx, .doc, .pdf)")
        return False

    print(f"找到 {len(file_list)} 个文件待处理")

    processed_count = 0
    for file_path in file_list:
        # Safely print filename by removing problematic characters
        safe_name = str(file_path.name).encode('ascii', errors='ignore').decode('ascii', errors='replace')
        if not safe_name.strip():
            safe_name = str(file_path.stem)[:20] + file_path.suffix  # fallback to basic name
        print(f"Processing: {safe_name[:30]}...")

        # Define output file path
        output_file = Path(output_dir) / f"{file_path.stem}.pdf"

        # Process based on file extension
        if file_path.suffix.lower() == '.pdf':
            # Copy PDF file directly to output
            success = copy_pdf_to_output(file_path, output_file)
        elif file_path.suffix.lower() == '.docx':
            # Convert DOCX to PDF
            success = convert_docx_to_pdf(file_path, output_file)
        elif file_path.suffix.lower() == '.doc':
            # Convert DOC to PDF
            success = convert_doc_to_pdf(file_path, output_file)
        else:
            safe_name = str(file_path.name).encode('ascii', errors='ignore').decode('ascii', errors='replace')
            if not safe_name.strip():
                safe_name = str(file_path.stem)[:20] + file_path.suffix
            print(f"Skipping unsupported file type: {safe_name[:30]}...")
            continue

        # Safely print results
        safe_input_name = str(file_path.name).encode('ascii', errors='ignore').decode('ascii', errors='replace')
        if not safe_input_name.strip():
            safe_input_name = str(file_path.stem)[:15] + file_path.suffix
        safe_output_name = str(output_file.name).encode('ascii', errors='ignore').decode('ascii', errors='replace')
        if not safe_output_name.strip():
            safe_output_name = str(output_file.stem)[:15] + output_file.suffix

        if success:
            print(f"[OK] Processed: {safe_input_name[:20]}... -> {safe_output_name[:20]}...")
            processed_count += 1
        else:
            print(f"[ERR] Failed: {safe_input_name[:30]}...")

    print(f"\nProcessing complete! Converted {processed_count} files")
    return processed_count > 0


def main():
    parser = argparse.ArgumentParser(description="批量将Word文档转换为PDF文件")
    parser.add_argument(
        "-i", "--input",
        dest="input_path",
        type=str,
        required=True,
        help="输入文件路径或目录，支持 .docx .doc .pdf 文件"
    )
    parser.add_argument(
        "-o", "--output",
        dest="output_dir",
        type=str,
        required=True,
        help="输出目录路径"
    )

    args = parser.parse_args()

    # Run batch conversion
    success = batch_convert_word_to_pdf(args.input_path, args.output_dir)

    if success:
        print("\nAll files have been successfully converted or copied to the output directory!")
        sys.exit(0)
    else:
        print("\nErrors occurred during processing!")
        sys.exit(1)


if __name__ == "__main__":
    # For direct execution with default paths
    if len(sys.argv) == 1:
        # 请根据您的实际路径修改以下两个变量
        input_path = r"C:\Users\Qzj\Desktop\projrct\MinerU\input_file_mix"  # 输入文件或目录路径
        output_dir = r"C:\Users\Qzj\Desktop\projrct\MinerU\input_file_pdf"  # 输出目录路径

        print("使用默认路径进行转换...")
        print(f"输入路径: {input_path}")
        print(f"输出路径: {output_dir}")

        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)

        # 运行转换
        success = batch_convert_word_to_pdf(input_path, output_dir)

        if success:
            print("\n所有文件已成功转换或复制至输出目录!")
        else:
            print("\n处理过程中出现错误!")
    else:
        main()
